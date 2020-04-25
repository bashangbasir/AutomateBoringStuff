import docx
import os 

os.chdir(r"F:\AutomateBoringStuff\WordExcelPDFDocument\Word")

#create new document
d = docx.Document()

#add paragraph to document
d.add_paragraph("First paragraph is created !")
#add paragraph to document with style
d.add_paragraph("Second paragraph is created !",style="Title")

# add run to created paragraph 
p1 = d.paragraphs[0]
p1.add_run("Add this using run")
#set run as bold 
p1.runs[0].bold = True
d.save("demo3.docx")
