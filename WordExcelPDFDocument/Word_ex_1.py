import docx
import os 

#change to word file working directory 
os.chdir(r"F:\AutomateBoringStuff\WordExcelPDFDocument\Word")

d = docx.Document("demo.docx") #return document object

p = d.paragraphs # return list of paragraphs object

# to get the text value of paragraph you can use index to spcefiy the paragraph and text attribute
p1_text = d.paragraphs[0].text # first index start at 0

# run() method is to check how many setting of run in chertain paragraph.
p2 = d.paragraphs[1]
runs = p2.runs # will return list of run object

#to check the run text
for run in runs:
    print(run.text)

# also can check what are the run format 
is_bold = p2.runs[2].bold
print("p2 runs[2] is bold? : ",is_bold)
is_italic = p2.runs[4].italic
print("p2 runs[4] is italic? : ",is_italic)

#to change run text and make it underline
p2.runs[4].text = "italic and underlined"
p2.runs[4].underlined = True

# to change style of paragraph 
p2_style = p2.style
print("p2 style is : ",p2_style)
p2.style = "Title"
p2_style = p2.style
print("p2 style is : ",p2_style)

d.save("demo2.docx")

